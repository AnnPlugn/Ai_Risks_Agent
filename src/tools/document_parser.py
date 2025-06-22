# src/tools/document_parser.py
"""
Парсер документов для извлечения информации об ИИ-агентах
Поддерживает Word (.docx), Excel (.xlsx), PDF и текстовые файлы
"""

import re
import json
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass
from datetime import datetime

# Импорты для парсинга документов
import docx
from docx.document import Document
import openpyxl
from openpyxl.workbook import Workbook
import PyPDF2
import pdfplumber

from ..utils.logger import get_logger, LogContext


@dataclass
class ParsedDocument:
    """Результат парсинга документа"""
    file_path: str
    file_type: str
    content: str
    metadata: Dict[str, Any]
    sections: Dict[str, str]
    tables: List[Dict[str, Any]]
    success: bool
    error_message: Optional[str] = None
    parsing_time: float = 0.0


class DocumentParserError(Exception):
    """Исключение при парсинге документов"""
    pass


class BaseDocumentParser:
    """Базовый класс для парсеров документов"""
    
    def __init__(self):
        self.logger = get_logger()
        self.supported_extensions = []
    
    def can_parse(self, file_path: Union[str, Path]) -> bool:
        """Проверка, может ли парсер обработать файл"""
        path = Path(file_path)
        return path.suffix.lower() in self.supported_extensions
    
    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """Базовый метод парсинга (переопределяется в наследниках)"""
        raise NotImplementedError
    
    def _extract_metadata(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Извлечение метаданных файла"""
        path = Path(file_path)
        stat = path.stat()
        
        return {
            "file_name": path.name,
            "file_size": stat.st_size,
            "created_time": datetime.fromtimestamp(stat.st_ctime),
            "modified_time": datetime.fromtimestamp(stat.st_mtime),
            "extension": path.suffix.lower()
        }


class WordDocumentParser(BaseDocumentParser):
    """Парсер для Word документов (.docx)"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']
    
    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """Парсинг Word документа"""
        start_time = datetime.now()
        path = Path(file_path)
        
        try:
            with LogContext("parse_word_document", "document_parser", "document_parser"):
                # Открываем документ
                doc = docx.Document(str(path))
                
                # Извлекаем текст по параграфам
                paragraphs = []
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        paragraphs.append(text)
                
                content = "\n".join(paragraphs)
                
                # Извлекаем заголовки и секции
                sections = self._extract_sections_from_word(doc)
                
                # Извлекаем таблицы
                tables = self._extract_tables_from_word(doc)
                
                # Метаданные документа
                metadata = self._extract_metadata(path)
                metadata.update(self._extract_word_properties(doc))
                
                execution_time = (datetime.now() - start_time).total_seconds()
                
                self.logger.log_document_parsing("document_parser", str(path), "Word", True)
                
                return ParsedDocument(
                    file_path=str(path),
                    file_type="word",
                    content=content,
                    metadata=metadata,
                    sections=sections,
                    tables=tables,
                    success=True,
                    parsing_time=execution_time
                )
                
        except Exception as e:
            execution_time = (datetime.now() - start_time).total_seconds()
            self.logger.log_document_parsing("document_parser", str(path), "Word", False)
            
            return ParsedDocument(
                file_path=str(path),
                file_type="word",
                content="",
                metadata=self._extract_metadata(path),
                sections={},
                tables=[],
                success=False,
                error_message=str(e),
                parsing_time=execution_time
            )
    
    def _extract_sections_from_word(self, doc: Document) -> Dict[str, str]:
        """Извлечение секций из Word документа по заголовкам"""
        sections = {}
        current_section = "introduction"
        current_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Проверяем, является ли параграф заголовком
            if self._is_heading(paragraph):
                # Сохраняем предыдущую секцию
                if current_content:
                    sections[current_section] = "\n".join(current_content)
                
                # Начинаем новую секцию
                current_section = self._normalize_section_name(text)
                current_content = []
            else:
                current_content.append(text)
        
        # Сохраняем последнюю секцию
        if current_content:
            sections[current_section] = "\n".join(current_content)
        
        return sections
    
    def _extract_tables_from_word(self, doc: Document) -> List[Dict[str, Any]]:
        """Извлечение таблиц из Word документа"""
        tables = []
        
        for i, table in enumerate(doc.tables):
            table_data = {
                "table_index": i,
                "rows": [],
                "headers": []
            }
            
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                
                if row_idx == 0:
                    table_data["headers"] = row_data
                else:
                    table_data["rows"].append(row_data)
            
            tables.append(table_data)
        
        return tables
    
    def _extract_word_properties(self, doc: Document) -> Dict[str, Any]:
        """Извлечение свойств Word документа"""
        props = {}
        
        try:
            core_props = doc.core_properties
            props.update({
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "created": core_props.created,
                "modified": core_props.modified
            })
        except Exception:
            pass
        
        return props
    
    def _is_heading(self, paragraph) -> bool:
        """Проверка, является ли параграф заголовком"""
        # Проверяем стиль параграфа
        if paragraph.style.name.startswith('Heading'):
            return True
        
        # Проверяем форматирование (жирный, размер шрифта)
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.bold and len(paragraph.text.strip()) < 100:
                return True
        
        # Проверяем паттерны заголовков
        text = paragraph.text.strip()
        heading_patterns = [
            r'^\d+\.\s+',  # "1. Заголовок"
            r'^[А-ЯA-Z][А-ЯA-Z\s]+:$',  # "ЗАГОЛОВОК:"
            r'^[А-Я][а-я\s]+:$'  # "Заголовок:"
        ]
        
        for pattern in heading_patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    def _normalize_section_name(self, heading_text: str) -> str:
        """Нормализация названия секции"""
        # Убираем номера и специальные символы
        normalized = re.sub(r'^\d+\.?\s*', '', heading_text)
        normalized = re.sub(r'[:\.]$', '', normalized)
        normalized = normalized.lower().strip()
        
        # Маппинг на стандартные названия секций
        section_mapping = {
            'общая информация': 'general_info',
            'общие сведения': 'general_info', 
            'описание': 'description',
            'назначение': 'purpose',
            'техническая спецификация': 'technical_spec',
            'технические характеристики': 'technical_spec',
            'архитектура': 'architecture',
            'промпты': 'prompts',
            'инструкции': 'instructions',
            'системные промпты': 'system_prompts',
            'ограничения': 'guardrails',
            'меры безопасности': 'guardrails',
            'бизнес-контекст': 'business_context',
            'целевая аудитория': 'target_audience',
            'риски': 'risks',
            'безопасность': 'security',
            'интеграции': 'integrations',
            'api': 'integrations',
            'данные': 'data_access',
            'персональные данные': 'data_access'
        }
        
        return section_mapping.get(normalized, normalized.replace(' ', '_'))


class ExcelDocumentParser(BaseDocumentParser):
    """Парсер для Excel документов (.xlsx)"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.xlsx', '.xls']
    
    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """Парсинг Excel документа"""
        start_time = datetime.now()
        path = Path(file_path)
        
        try:
            with LogContext("parse_excel_document", "document_parser", "document_parser"):
                # Открываем Excel файл
                workbook = openpyxl.load_workbook(str(path), data_only=True)
                
                content_parts = []
                sections = {}
                tables = []
                
                # Обрабатываем каждый лист
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    
                    # Извлекаем данные из листа
                    sheet_content, sheet_tables = self._extract_sheet_data(sheet, sheet_name)
                    
                    content_parts.append(f"=== Лист: {sheet_name} ===\n{sheet_content}")
                    sections[f"sheet_{sheet_name.lower()}"] = sheet_content
                    tables.extend(sheet_tables)
                
                content = "\n\n".join(content_parts)
                
                # Метаданные
                metadata = self._extract_metadata(path)
                metadata.update(self._extract_excel_properties(workbook))
                
                execution_time = (datetime.now() - start_time).total_seconds()
                
                self.logger.log_document_parsing("document_parser", str(path), "Excel", True)
                
                return ParsedDocument(
                    file_path=str(path),
                    file_type="excel",
                    content=content,
                    metadata=metadata,
                    sections=sections,
                    tables=tables,
                    success=True,
                    parsing_time=execution_time
                )
                
        except Exception as e:
            execution_time = (datetime.now() - start_time).total_seconds()
            self.logger.log_document_parsing("document_parser", str(path), "Excel", False)
            
            return ParsedDocument(
                file_path=str(path),
                file_type="excel",
                content="",
                metadata=self._extract_metadata(path),
                sections={},
                tables=[],
                success=False,
                error_message=str(e),
                parsing_time=execution_time
            )
    
    def _extract_sheet_data(self, sheet, sheet_name: str) -> tuple[str, List[Dict[str, Any]]]:
        """Извлечение данных из листа Excel"""
        content_lines = []
        tables = []
        
        # Находим границы данных
        max_row = sheet.max_row
        max_col = sheet.max_column
        
        if max_row == 1 and max_col == 1:
            return "", []
        
        # Извлекаем данные построчно
        current_table = []
        headers = []
        
        for row_idx, row in enumerate(sheet.iter_rows(max_row=max_row, max_col=max_col), 1):
            row_data = []
            has_data = False
            
            for cell in row:
                value = cell.value
                if value is not None:
                    value = str(value).strip()
                    has_data = True
                else:
                    value = ""
                row_data.append(value)
            
            if has_data:
                # Первая строка с данными - заголовки
                if not headers and any(row_data):
                    headers = row_data
                    current_table.append(row_data)
                else:
                    current_table.append(row_data)
                
                # Добавляем строку в контент
                content_lines.append(" | ".join(filter(None, row_data)))
        
        # Создаем таблицу если есть данные
        if current_table:
            tables.append({
                "sheet_name": sheet_name,
                "headers": headers if headers else [],
                "rows": current_table[1:] if headers else current_table,
                "total_rows": len(current_table),
                "total_cols": len(headers) if headers else max_col
            })
        
        content = "\n".join(content_lines)
        return content, tables
    
    def _extract_excel_properties(self, workbook: Workbook) -> Dict[str, Any]:
        """Извлечение свойств Excel документа"""
        props = {
            "sheet_names": workbook.sheetnames,
            "sheet_count": len(workbook.sheetnames)
        }
        
        try:
            if hasattr(workbook, 'properties'):
                core_props = workbook.properties
                props.update({
                    "title": getattr(core_props, 'title', ''),
                    "author": getattr(core_props, 'creator', ''),
                    "subject": getattr(core_props, 'subject', ''),
                    "created": getattr(core_props, 'created', None),
                    "modified": getattr(core_props, 'modified', None)
                })
        except Exception:
            pass
        
        return props


class PDFDocumentParser(BaseDocumentParser):
    """Парсер для PDF документов"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.pdf']
    
    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """Парсинг PDF документа"""
        start_time = datetime.now()
        path = Path(file_path)
        
        try:
            with LogContext("parse_pdf_document", "document_parser", "document_parser"):
                # Пробуем pdfplumber для лучшего извлечения текста
                content_parts = []
                tables = []
                
                with pdfplumber.open(str(path)) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        # Извлекаем текст
                        text = page.extract_text()
                        if text:
                            content_parts.append(f"=== Страница {page_num} ===\n{text}")
                        
                        # Пытаемся извлечь таблицы
                        page_tables = page.extract_tables()
                        for table_idx, table in enumerate(page_tables):
                            if table:
                                tables.append({
                                    "page": page_num,
                                    "table_index": table_idx,
                                    "headers": table[0] if table else [],
                                    "rows": table[1:] if len(table) > 1 else [],
                                    "raw_data": table
                                })
                
                content = "\n\n".join(content_parts)
                
                # Извлекаем секции по заголовкам
                sections = self._extract_sections_from_text(content)
                
                # Метаданные
                metadata = self._extract_metadata(path)
                metadata.update(self._extract_pdf_properties(str(path)))
                
                execution_time = (datetime.now() - start_time).total_seconds()
                
                self.logger.log_document_parsing("document_parser", str(path), "PDF", True)
                
                return ParsedDocument(
                    file_path=str(path),
                    file_type="pdf",
                    content=content,
                    metadata=metadata,
                    sections=sections,
                    tables=tables,
                    success=True,
                    parsing_time=execution_time
                )
                
        except Exception as e:
            execution_time = (datetime.now() - start_time).total_seconds()
            self.logger.log_document_parsing("document_parser", str(path), "PDF", False)
            
            return ParsedDocument(
                file_path=str(path),
                file_type="pdf",
                content="",
                metadata=self._extract_metadata(path),
                sections={},
                tables=[],
                success=False,
                error_message=str(e),
                parsing_time=execution_time
            )
    
    def _extract_sections_from_text(self, text: str) -> Dict[str, str]:
        """Извлечение секций из текста PDF по заголовкам"""
        sections = {}
        lines = text.split('\n')
        
        current_section = "introduction"
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Проверяем, является ли строка заголовком
            if self._is_text_heading(line):
                # Сохраняем предыдущую секцию
                if current_content:
                    sections[current_section] = "\n".join(current_content)
                
                # Начинаем новую секцию
                current_section = self._normalize_section_name(line)
                current_content = []
            else:
                current_content.append(line)
        
        # Сохраняем последнюю секцию
        if current_content:
            sections[current_section] = "\n".join(current_content)
        
        return sections
    
    def _extract_pdf_properties(self, file_path: str) -> Dict[str, Any]:
        """Извлечение свойств PDF документа"""
        props = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                props.update({
                    "page_count": len(pdf_reader.pages),
                    "encrypted": pdf_reader.is_encrypted
                })
                
                # Метаданные PDF
                if pdf_reader.metadata:
                    metadata = pdf_reader.metadata
                    props.update({
                        "title": metadata.get('/Title', ''),
                        "author": metadata.get('/Author', ''),
                        "subject": metadata.get('/Subject', ''),
                        "creator": metadata.get('/Creator', ''),
                        "producer": metadata.get('/Producer', ''),
                        "creation_date": metadata.get('/CreationDate', ''),
                        "modification_date": metadata.get('/ModDate', '')
                    })
        except Exception:
            pass
        
        return props
    
    def _is_text_heading(self, line: str) -> bool:
        """Проверка, является ли строка заголовком в тексте"""
        # Паттерны заголовков
        heading_patterns = [
            r'^\d+\.\s+',  # "1. Заголовок"
            r'^[А-ЯA-Z][А-ЯA-Z\s]+$',  # "ЗАГОЛОВОК"
            r'^[А-Я][а-я\s]+:$',  # "Заголовок:"
            r'^={3,}\s+.+\s+={3,}$'  # "=== Заголовок ==="
        ]
        
        for pattern in heading_patterns:
            if re.match(pattern, line):
                return True
        
        # Проверяем короткие строки в верхнем регистре
        if len(line) < 100 and line.isupper() and len(line.split()) > 1:
            return True
        
        return False


class TextDocumentParser(BaseDocumentParser):
    """Парсер для текстовых файлов (.txt, .md, .py, .js и т.д.)"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.txt', '.md', '.py', '.js', '.json', '.yaml', '.yml']
    
    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """Парсинг текстового файла"""
        start_time = datetime.now()
        path = Path(file_path)
        
        try:
            with LogContext("parse_text_document", "document_parser", "document_parser"):
                # Определяем кодировку
                encoding = self._detect_encoding(path)
                
                # Читаем файл
                with open(path, 'r', encoding=encoding) as file:
                    content = file.read()
                
                # Извлекаем секции в зависимости от типа файла
                sections = self._extract_sections_by_file_type(content, path.suffix.lower())
                
                # Метаданные
                metadata = self._extract_metadata(path)
                metadata.update({
                    "encoding": encoding,
                    "line_count": len(content.splitlines()),
                    "char_count": len(content)
                })
                
                execution_time = (datetime.now() - start_time).total_seconds()
                
                self.logger.log_document_parsing("document_parser", str(path), f"Text({path.suffix})", True)
                
                return ParsedDocument(
                    file_path=str(path),
                    file_type=f"text_{path.suffix[1:]}",
                    content=content,
                    metadata=metadata,
                    sections=sections,
                    tables=[],
                    success=True,
                    parsing_time=execution_time
                )
                
        except Exception as e:
            execution_time = (datetime.now() - start_time).total_seconds()
            self.logger.log_document_parsing("document_parser", str(path), f"Text({path.suffix})", False)
            
            return ParsedDocument(
                file_path=str(path),
                file_type=f"text_{path.suffix[1:]}",
                content="",
                metadata=self._extract_metadata(path),
                sections={},
                tables=[],
                success=False,
                error_message=str(e),
                parsing_time=execution_time
            )
    
    def _detect_encoding(self, path: Path) -> str:
        """Определение кодировки файла"""
        # Список кодировок для проверки
        encodings = ['utf-8', 'cp1251', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as file:
                    file.read()
                return encoding
            except UnicodeDecodeError:
                continue
        
        return 'utf-8'  # По умолчанию
    
    def _extract_sections_by_file_type(self, content: str, extension: str) -> Dict[str, str]:
        """Извлечение секций в зависимости от типа файла"""
        sections = {}
        
        if extension == '.md':
            sections = self._extract_markdown_sections(content)
        elif extension in ['.py', '.js']:
            sections = self._extract_code_sections(content, extension)
        elif extension in ['.json', '.yaml', '.yml']:
            sections = self._extract_config_sections(content, extension)
        else:
            # Для обычных текстовых файлов
            sections = {"content": content}
        
        return sections
    
    def _extract_markdown_sections(self, content: str) -> Dict[str, str]:
        """Извлечение секций из Markdown"""
        sections = {}
        lines = content.split('\n')
        
        current_section = "introduction"
        current_content = []
        
        for line in lines:
            # Проверяем заголовки Markdown
            if line.startswith('#'):
                # Сохраняем предыдущую секцию
                if current_content:
                    sections[current_section] = "\n".join(current_content)
                
                # Начинаем новую секцию
                heading = line.lstrip('#').strip()
                current_section = self._normalize_section_name(heading)
                current_content = []
            else:
                current_content.append(line)
        
        # Сохраняем последнюю секцию
        if current_content:
            sections[current_section] = "\n".join(current_content)
        
        return sections
    
    def _extract_code_sections(self, content: str, extension: str) -> Dict[str, str]:
        """Извлечение секций из кода"""
        sections = {
            "full_code": content,
            "comments": self._extract_comments(content, extension),
            "functions": self._extract_functions(content, extension),
            "classes": self._extract_classes(content, extension)
        }
        
        return sections
    
    def _extract_config_sections(self, content: str, extension: str) -> Dict[str, str]:
        """Извлечение секций из конфигурационных файлов"""
        sections = {"config": content}
        
        try:
            if extension == '.json':
                data = json.loads(content)
                sections["parsed_json"] = json.dumps(data, indent=2, ensure_ascii=False)
            # Для YAML можно добавить парсинг если нужно
        except Exception:
            pass
        
        return sections
    
    def _extract_comments(self, content: str, extension: str) -> str:
        """Извлечение комментариев из кода"""
        comments = []
        lines = content.split('\n')
        
        comment_patterns = {
            '.py': [r'^\s*#', r'""".*?"""', r"'''.*?'''"],
            '.js': [r'^\s*//', r'/\*.*?\*/']
        }
        
        patterns = comment_patterns.get(extension, [])
        
        for line in lines:
            for pattern in patterns:
                if re.match(pattern, line.strip()):
                    comments.append(line.strip())
                    break
        
        return "\n".join(comments)
    
    def _extract_functions(self, content: str, extension: str) -> str:
        """Извлечение функций из кода"""
        functions = []
        
        function_patterns = {
            '.py': r'^\s*def\s+(\w+)\s*\([^)]*\):',
            '.js': r'^\s*function\s+(\w+)\s*\([^)]*\)\s*{|^\s*const\s+(\w+)\s*=\s*\([^)]*\)\s*=>'
        }
        
        pattern = function_patterns.get(extension)
        if pattern:
            for match in re.finditer(pattern, content, re.MULTILINE):
                functions.append(match.group(0))
        
        return "\n".join(functions)
    
    def _extract_classes(self, content: str, extension: str) -> str:
        """Извлечение классов из кода"""
        classes = []
        
        class_patterns = {
            '.py': r'^\s*class\s+(\w+).*?:',
            '.js': r'^\s*class\s+(\w+).*?{'
        }
        
        pattern = class_patterns.get(extension)
        if pattern:
            for match in re.finditer(pattern, content, re.MULTILINE):
                classes.append(match.group(0))
        
        return "\n".join(classes)


# ===============================
# Главный парсер документов
# ===============================

class DocumentParser:
    """Главный класс для парсинга различных типов документов"""
    
    def __init__(self):
        self.parsers = [
            WordDocumentParser(),
            ExcelDocumentParser(),
            PDFDocumentParser(),
            TextDocumentParser()
        ]
        self.logger = get_logger()
    
    def can_parse(self, file_path: Union[str, Path]) -> bool:
        """Проверка, может ли система парсить файл"""
        return any(parser.can_parse(file_path) for parser in self.parsers)
    
    def parse_document(self, file_path: Union[str, Path]) -> ParsedDocument:
        """Парсинг документа с автоматическим выбором парсера"""
        path = Path(file_path)
        
        if not path.exists():
            return ParsedDocument(
                file_path=str(path),
                file_type="unknown",
                content="",
                metadata={},
                sections={},
                tables=[],
                success=False,
                error_message=f"Файл не найден: {path}"
            )
        
        # Ищем подходящий парсер
        for parser in self.parsers:
            if parser.can_parse(path):
                bound_logger = self.logger.bind_context("document_parser", "document_parser")
                bound_logger.info(f"📄 Парсинг файла: {path.name} ({parser.__class__.__name__})")
                
                result = parser.parse(path)
                
                if result.success:
                    bound_logger.info(f"✅ Файл распарсен: {len(result.content)} символов, {len(result.sections)} секций")
                else:
                    bound_logger.error(f"❌ Ошибка парсинга: {result.error_message}")
                
                return result
        
        # Если парсер не найден
        return ParsedDocument(
            file_path=str(path),
            file_type="unsupported",
            content="",
            metadata=self._extract_basic_metadata(path),
            sections={},
            tables=[],
            success=False,
            error_message=f"Неподдерживаемый тип файла: {path.suffix}"
        )
    
    def parse_multiple_documents(self, file_paths: List[Union[str, Path]]) -> List[ParsedDocument]:
        """Парсинг нескольких документов"""
        results = []
        
        for file_path in file_paths:
            result = self.parse_document(file_path)
            results.append(result)
        
        # Логируем итоговую статистику
        successful = sum(1 for r in results if r.success)
        failed = len(results) - successful
        
        bound_logger = self.logger.bind_context("document_parser", "document_parser")
        bound_logger.info(f"📊 Парсинг завершен: {successful} успешно, {failed} ошибок")
        
        return results
    
    def extract_agent_info_from_documents(self, documents: List[ParsedDocument]) -> Dict[str, Any]:
        """
        Извлечение информации об агенте из распарсенных документов
        Специализированный метод для агентов ИИ
        """
        agent_info = {
            "name": "",
            "description": "",
            "technical_specs": {},
            "prompts": [],
            "guardrails": [],
            "business_context": {},
            "source_files": []
        }
        
        for doc in documents:
            if not doc.success:
                continue
            
            agent_info["source_files"].append(doc.file_path)
            
            # Извлекаем информацию из секций
            for section_name, section_content in doc.sections.items():
                self._extract_info_from_section(agent_info, section_name, section_content)
            
            # Извлекаем информацию из таблиц
            for table in doc.tables:
                self._extract_info_from_table(agent_info, table)
            
            # Извлекаем информацию из метаданных
            if "title" in doc.metadata and doc.metadata["title"]:
                if not agent_info["name"]:
                    agent_info["name"] = doc.metadata["title"]
        
        return agent_info
    
    def _extract_info_from_section(self, agent_info: Dict[str, Any], section_name: str, content: str):
        """Извлечение информации из секции документа"""
        content = content.strip()
        if not content:
            return
        
        # Маппинг секций на поля агента
        if section_name in ["general_info", "description", "introduction"]:
            if not agent_info["description"]:
                agent_info["description"] = content
        
        elif section_name in ["technical_spec", "architecture"]:
            agent_info["technical_specs"][section_name] = content
        
        elif section_name in ["prompts", "system_prompts", "instructions"]:
            # Извлекаем промпты из текста
            prompts = self._extract_prompts_from_text(content)
            agent_info["prompts"].extend(prompts)
        
        elif section_name in ["guardrails", "security", "limitations"]:
            # Извлекаем ограничения
            guardrails = self._extract_guardrails_from_text(content)
            agent_info["guardrails"].extend(guardrails)
        
        elif section_name in ["business_context", "target_audience"]:
            agent_info["business_context"][section_name] = content
    
    def _extract_info_from_table(self, agent_info: Dict[str, Any], table: Dict[str, Any]):
        """Извлечение информации из таблицы"""
        if not table.get("headers") or not table.get("rows"):
            return
        
        headers = [h.lower().strip() if h else "" for h in table["headers"]]
        
        # Ищем таблицы с техническими характеристиками
        tech_indicators = ["параметр", "характеристика", "значение", "описание"]
        if any(indicator in " ".join(headers) for indicator in tech_indicators):
            for row in table["rows"]:
                if len(row) >= 2 and row[0] and row[1]:
                    key = str(row[0]).strip()
                    value = str(row[1]).strip()
                    agent_info["technical_specs"][key] = value
    
    def _extract_prompts_from_text(self, text: str) -> List[str]:
        """Извлечение промптов из текста"""
        prompts = []
        
        # Паттерны для поиска промптов
        prompt_patterns = [
            r'Системный промпт[:\s]*(.+?)(?=\n\n|\n[А-Я]|$)',
            r'Промпт[:\s]*(.+?)(?=\n\n|\n[А-Я]|$)',
            r'Инструкция[:\s]*(.+?)(?=\n\n|\n[А-Я]|$)',
            r'"([^"]+)"',  # Текст в кавычках
            r'```\s*(.+?)\s*```'  # Код блоки
        ]
        
        for pattern in prompt_patterns:
            matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
            for match in matches:
                cleaned = match.strip()
                if len(cleaned) > 10:  # Минимальная длина промпта
                    prompts.append(cleaned)
        
        # Разбиваем по строкам если нет специальных паттернов
        if not prompts:
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if len(line) > 20 and not line.startswith('•') and not line.startswith('-'):
                    prompts.append(line)
        
        return prompts
    
    def _extract_guardrails_from_text(self, text: str) -> List[str]:
        """Извлечение ограничений безопасности из текста"""
        guardrails = []
        
        # Паттерны для ограничений
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            
            # Ищем строки с ограничениями
            restriction_indicators = [
                'не должен', 'запрещено', 'нельзя', 'ограничение',
                'не разрешается', 'недопустимо', 'исключить'
            ]
            
            if any(indicator in line.lower() for indicator in restriction_indicators):
                if len(line) > 10:
                    guardrails.append(line)
            
            # Ищем списки ограничений
            if line.startswith(('•', '-', '*')) and len(line) > 10:
                guardrails.append(line[1:].strip())
        
        return guardrails
    
    def _extract_basic_metadata(self, path: Path) -> Dict[str, Any]:
        """Извлечение базовых метаданных файла"""
        try:
            stat = path.stat()
            return {
                "file_name": path.name,
                "file_size": stat.st_size,
                "extension": path.suffix.lower(),
                "created_time": datetime.fromtimestamp(stat.st_ctime),
                "modified_time": datetime.fromtimestamp(stat.st_mtime)
            }
        except Exception:
            return {"file_name": path.name, "extension": path.suffix.lower()}
    
    def get_supported_extensions(self) -> List[str]:
        """Получение списка поддерживаемых расширений"""
        extensions = []
        for parser in self.parsers:
            extensions.extend(parser.supported_extensions)
        return sorted(list(set(extensions)))
    
    def get_parsing_stats(self, documents: List[ParsedDocument]) -> Dict[str, Any]:
        """Получение статистики парсинга"""
        if not documents:
            return {}
        
        successful = [d for d in documents if d.success]
        failed = [d for d in documents if not d.success]
        
        stats = {
            "total_documents": len(documents),
            "successful": len(successful),
            "failed": len(failed),
            "success_rate": len(successful) / len(documents) * 100,
            "total_content_length": sum(len(d.content) for d in successful),
            "total_sections": sum(len(d.sections) for d in successful),
            "total_tables": sum(len(d.tables) for d in successful),
            "avg_parsing_time": sum(d.parsing_time for d in documents) / len(documents),
            "file_types": {}
        }
        
        # Статистика по типам файлов
        for doc in documents:
            file_type = doc.file_type
            if file_type not in stats["file_types"]:
                stats["file_types"][file_type] = {"count": 0, "successful": 0}
            
            stats["file_types"][file_type]["count"] += 1
            if doc.success:
                stats["file_types"][file_type]["successful"] += 1
        
        return stats


# ===============================
# Утилитарные функции
# ===============================

def create_document_parser() -> DocumentParser:
    """Фабрика для создания парсера документов"""
    return DocumentParser()


def parse_agent_documents(
    file_paths: List[Union[str, Path]], 
    extract_agent_info: bool = True
) -> tuple[List[ParsedDocument], Optional[Dict[str, Any]]]:
    """
    Удобная функция для парсинга документов агента
    
    Args:
        file_paths: Список путей к файлам
        extract_agent_info: Извлекать ли информацию об агенте
    
    Returns:
        Tuple из (результаты парсинга, информация об агенте)
    """
    parser = create_document_parser()
    
    # Парсим документы
    parsed_docs = parser.parse_multiple_documents(file_paths)
    
    # Извлекаем информацию об агенте если нужно
    agent_info = None
    if extract_agent_info:
        agent_info = parser.extract_agent_info_from_documents(parsed_docs)
    
    return parsed_docs, agent_info


def get_document_summary(parsed_doc: ParsedDocument) -> Dict[str, Any]:
    """Получение краткой сводки о документе"""
    return {
        "file_name": Path(parsed_doc.file_path).name,
        "file_type": parsed_doc.file_type,
        "success": parsed_doc.success,
        "content_length": len(parsed_doc.content),
        "sections_count": len(parsed_doc.sections),
        "tables_count": len(parsed_doc.tables),
        "parsing_time": parsed_doc.parsing_time,
        "error": parsed_doc.error_message if not parsed_doc.success else None
    }


# Экспорт основных классов и функций
__all__ = [
    "DocumentParser",
    "ParsedDocument", 
    "DocumentParserError",
    "WordDocumentParser",
    "ExcelDocumentParser", 
    "PDFDocumentParser",
    "TextDocumentParser",
    "create_document_parser",
    "parse_agent_documents",
    "get_document_summary"
]