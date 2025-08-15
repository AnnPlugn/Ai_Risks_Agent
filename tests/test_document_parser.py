# test_document_parser.py
"""
Тестовый скрипт для диагностики и исправления проблем с парсером документов
"""

import sys
from pathlib import Path
import tempfile
import json
from typing import List, Dict, Any

# Добавляем путь к src
sys.path.insert(0, str(Path(__file__).parent / "src"))

def check_dependencies():
    """Проверка необходимых зависимостей для парсинга документов"""
    print("🔍 Проверка зависимостей для парсинга документов...")
    
    dependencies = {
        'docx': 'python-docx',
        'openpyxl': 'openpyxl', 
        'PyPDF2': 'PyPDF2',
        'pdfplumber': 'pdfplumber'
    }
    
    missing_deps = []
    
    for module, package in dependencies.items():
        try:
            __import__(module)
            print(f"✅ {module} ({package}) - установлен")
        except ImportError:
            print(f"❌ {module} ({package}) - НЕ УСТАНОВЛЕН")
            missing_deps.append(package)
    
    if missing_deps:
        print(f"\n🚨 Недостающие зависимости:")
        print(f"Установите их командой: pip install {' '.join(missing_deps)}")
        return False
    
    print("✅ Все зависимости установлены")
    return True

def create_test_documents() -> Dict[str, Path]:
    """Создание тестовых документов для проверки парсинга"""
    print("\n📝 Создание тестовых документов...")
    
    test_docs = {}
    temp_dir = Path(tempfile.mkdtemp())
    
    # 1. Создаем тестовый текстовый файл
    txt_file = temp_dir / "test_agent.txt"
    txt_content = """Агент: TestAgent

Описание:
Тестовый ИИ-агент для банковских операций.

Функции:
- Обработка запросов клиентов
- Анализ рисков
- Предоставление рекомендаций

Ограничения:
- Не обрабатывает персональные данные без разрешения
- Требует подтверждения для финансовых операций
"""
    txt_file.write_text(txt_content, encoding='utf-8')
    test_docs['txt'] = txt_file
    
    # 2. Создаем тестовый JSON файл
    json_file = temp_dir / "agent_config.json"
    json_content = {
        "agent_name": "TestAgent",
        "model": "gpt-4",
        "temperature": 0.1,
        "system_prompt": "Ты - банковский ассистент",
        "guardrails": [
            "Не разглашай конфиденциальную информацию",
            "Всегда проверяй данные клиента"
        ]
    }
    json_file.write_text(json.dumps(json_content, ensure_ascii=False, indent=2), encoding='utf-8')
    test_docs['json'] = json_file
    
    # 3. Создаем тестовый Word документ (если доступен python-docx)
    try:
        import docx
        doc = docx.Document()
        doc.add_heading('Техническая спецификация TestAgent', 0)
        doc.add_heading('Общие сведения', level=1)
        p = doc.add_paragraph('TestAgent - это ИИ-агент для банковской сферы.')
        
        doc.add_heading('Технические характеристики', level=1)
        doc.add_paragraph('Модель: GPT-4')
        doc.add_paragraph('Температура: 0.1')
        doc.add_paragraph('Максимальное число токенов: 4096')
        
        doc.add_heading('Ограничения безопасности', level=1)
        doc.add_paragraph('1. Не обрабатывает персональные данные без согласия')
        doc.add_paragraph('2. Требует подтверждения для операций свыше 100,000 руб.')
        
        docx_file = temp_dir / "test_agent_spec.docx"
        doc.save(str(docx_file))
        test_docs['docx'] = docx_file
        print("✅ Word документ создан")
        
    except ImportError:
        print("⚠️ python-docx недоступен - Word документ не создан")
    
    # 4. Создаем тестовый Excel файл (если доступен openpyxl)
    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Агентные данные"
        
        # Заголовки
        ws['A1'] = 'Параметр'
        ws['B1'] = 'Значение'
        ws['C1'] = 'Описание'
        
        # Данные
        data = [
            ['Имя агента', 'TestAgent', 'Уникальное имя агента'],
            ['Модель', 'GPT-4', 'Используемая языковая модель'],
            ['Температура', '0.1', 'Параметр креативности'],
            ['Макс токены', '4096', 'Максимальная длина ответа'],
            ['Категория риска', 'Средний', 'Оценка операционного риска']
        ]
        
        for row, (param, value, desc) in enumerate(data, 2):
            ws[f'A{row}'] = param
            ws[f'B{row}'] = value
            ws[f'C{row}'] = desc
        
        xlsx_file = temp_dir / "agent_params.xlsx"
        wb.save(str(xlsx_file))
        test_docs['xlsx'] = xlsx_file
        print("✅ Excel файл создан")
        
    except ImportError:
        print("⚠️ openpyxl недоступен - Excel файл не создан")
    
    print(f"📁 Тестовые документы созданы в: {temp_dir}")
    return test_docs

def test_document_parser(test_docs: Dict[str, Path]):
    """Тестирование парсера документов"""
    print("\n🧪 Тестирование парсера документов...")
    
    try:
        from src.tools.document_parser import create_document_parser
        
        parser = create_document_parser()
        
        results = []
        
        for doc_type, file_path in test_docs.items():
            print(f"\n📄 Тестируется: {file_path.name} (тип: {doc_type})")
            
            # Проверяем, может ли парсер обработать файл
            can_parse = parser.can_parse(file_path)
            print(f"   Может парсить: {can_parse}")
            
            if can_parse:
                try:
                    # Пытаемся парсить документ
                    result = parser.parse_document(file_path)
                    
                    print(f"   Статус: {'✅ Успешно' if result.success else '❌ Ошибка'}")
                    
                    if result.success:
                        print(f"   Содержимое: {len(result.content)} символов")
                        print(f"   Секции: {len(result.sections)}")
                        print(f"   Таблицы: {len(result.tables)}")
                        print(f"   Время парсинга: {result.parsing_time:.2f}с")
                        
                        # Показываем первые 200 символов содержимого
                        preview = result.content[:200] + "..." if len(result.content) > 200 else result.content
                        print(f"   Превью: {preview}")
                        
                    else:
                        print(f"   Ошибка: {result.error_message}")
                    
                    results.append({
                        'file_type': doc_type,
                        'file_name': file_path.name,
                        'success': result.success,
                        'error': result.error_message if not result.success else None,
                        'content_length': len(result.content) if result.success else 0,
                        'sections_count': len(result.sections) if result.success else 0,
                        'parsing_time': result.parsing_time
                    })
                    
                except Exception as e:
                    print(f"   ❌ Исключение при парсинге: {e}")
                    results.append({
                        'file_type': doc_type,
                        'file_name': file_path.name,
                        'success': False,
                        'error': str(e),
                        'content_length': 0,
                        'sections_count': 0,
                        'parsing_time': 0
                    })
            else:
                print(f"   ⚠️ Парсер не поддерживает данный тип файла")
                results.append({
                    'file_type': doc_type,
                    'file_name': file_path.name,
                    'success': False,
                    'error': 'Unsupported file type',
                    'content_length': 0,
                    'sections_count': 0,
                    'parsing_time': 0
                })
        
        return results
        
    except Exception as e:
        print(f"❌ Критическая ошибка при тестировании парсера: {e}")
        import traceback
        traceback.print_exc()
        return []

def test_parser_components():
    """Детальная диагностика компонентов парсера"""
    print("\n🔧 Диагностика компонентов парсера...")
    
    try:
        from src.tools.document_parser import DocumentParser
        
        parser = DocumentParser()
        print(f"✅ Основной класс DocumentParser создан")
        print(f"   Количество парсеров: {len(parser.parsers)}")
        
        for i, p in enumerate(parser.parsers):
            parser_name = p.__class__.__name__
            supported_ext = getattr(p, 'supported_extensions', [])
            print(f"   {i+1}. {parser_name}: {supported_ext}")
            
            # Тестируем каждый парсер на совместимость
            test_extensions = ['.txt', '.py', '.docx', '.xlsx', '.pdf', '.json']
            for ext in test_extensions:
                if ext in supported_ext:
                    print(f"      ✅ Поддерживает {ext}")
                    
    except Exception as e:
        print(f"❌ Ошибка диагностики: {e}")
        import traceback
        traceback.print_exc()

def fix_parser_issues():
    """Попытка исправления выявленных проблем"""
    print("\n🛠️ Попытка исправления проблем...")
    
    # Проверим импорты в модуле парсера
    try:
        print("🔍 Проверка импортов в document_parser.py...")
        
        # Проверяем каждую библиотеку отдельно
        libraries = ['docx', 'openpyxl', 'PyPDF2', 'pdfplumber']
        
        for lib in libraries:
            try:
                module = __import__(lib)
                print(f"✅ {lib}: Доступен (версия: {getattr(module, '__version__', 'неизвестна')})")
            except ImportError as e:
                print(f"❌ {lib}: Недоступен - {e}")
                print(f"   Установите: pip install {lib}")
        
        # Пытаемся импортировать классы парсера
        from src.tools.document_parser import (
            WordDocumentParser, 
            ExcelDocumentParser, 
            PDFDocumentParser, 
            TextDocumentParser
        )
        print("✅ Все классы парсеров импортированы успешно")
        
        return True
        
    except Exception as e:
        print(f"❌ Ошибка при проверке парсера: {e}")
        import traceback
        traceback.print_exc()
        return False

def generate_test_report(results: List[Dict[str, Any]]):
    """Генерация отчета о тестировании"""
    print("\n📊 ОТЧЕТ О ТЕСТИРОВАНИИ ПАРСЕРА")
    print("=" * 50)
    
    total_tests = len(results)
    successful_tests = sum(1 for r in results if r['success'])
    failed_tests = total_tests - successful_tests
    
    print(f"Всего тестов: {total_tests}")
    print(f"Успешных: {successful_tests}")
    print(f"Неудачных: {failed_tests}")
    print(f"Процент успеха: {successful_tests/total_tests*100:.1f}%" if total_tests > 0 else "N/A")
    
    print("\nДетальные результаты:")
    for result in results:
        status = "✅" if result['success'] else "❌"
        print(f"{status} {result['file_type'].upper()}: {result['file_name']}")
        if not result['success']:
            print(f"    Ошибка: {result['error']}")
        else:
            print(f"    Содержимое: {result['content_length']} символов, {result['sections_count']} секций")
    
    # Рекомендации по исправлению
    print("\n💡 РЕКОМЕНДАЦИИ:")
    
    failed_types = [r['file_type'] for r in results if not r['success']]
    
    if 'docx' in failed_types:
        print("- Для Word файлов: pip install python-docx")
    if 'xlsx' in failed_types:
        print("- Для Excel файлов: pip install openpyxl")
    if 'pdf' in failed_types:
        print("- Для PDF файлов: pip install PyPDF2 pdfplumber")
    
    if failed_tests == 0:
        print("🎉 Все парсеры работают корректно!")

def main():
    """Основная функция диагностики парсера"""
    print("🚀 Диагностика и тестирование парсера документов")
    print("=" * 60)
    
    # 1. Проверяем зависимости
    deps_ok = check_dependencies()
    
    # 2. Диагностируем компоненты
    test_parser_components()
    
    # 3. Пытаемся исправить проблемы
    fix_parser_issues()
    
    # 4. Создаем тестовые документы
    test_docs = create_test_documents()
    
    # 5. Тестируем парсер
    if test_docs:
        results = test_document_parser(test_docs)
        
        # 6. Генерируем отчет
        if results:
            generate_test_report(results)
    
    print("\n" + "=" * 60)
    print("🏁 Диагностика завершена")
    
    if not deps_ok:
        print("\n🚨 КРИТИЧНО: Установите недостающие зависимости для корректной работы парсера!")
        return False
    
    return True

if __name__ == "__main__":
    main()